"""Tests for document generation (generate_docx module)."""
import os
import sys

import pytest
from docx import Document

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'tools'))
import generate_docx


@pytest.fixture
def sample_data():
    return {
        'AGREEMENT_ID': 'CAC-20260428-RN99900001',
        'AGREEMENT_NAME_DESCRIPTION': 'Test Agreement Description',
        'PREVIOUS_AGREEMENT': 'RN-PREV-001',
        'START_DATE_YYYY_MM_DD': '2026-04-01',
        'END_DATE_YYYY_MM_DD': '2027-03-31',
        'ITS_SERVICE': 'PROCUREMENT-SOFTWARE',
        'ITS_SERVICE_TYPE': 'SOFTWARE-ORACLE',
        'AGREEMENT_TYPE': 'Net New',
        'MONTH_BILLED': 'Apr',
        'MONTHLY_RECURRING': '5000',
        'IFIS_CODE': '100-123456-7890-ABCDEF-1234',
        'SERVICE_OWNER': 'DCO Business Services',
        'AGREEMENT_AUTHOR': 'Max Keeping',
        'COMMENTS': 'Test comment for validation',
    }


@pytest.fixture
def generated_doc(sample_data):
    """Generate a doc and yield the Document object; clean up temp file after."""
    path = generate_docx.generate(sample_data)
    doc = Document(path)
    yield doc
    os.unlink(path)


class TestDocGeneration:
    def test_generates_valid_docx(self, sample_data):
        path = generate_docx.generate(sample_data)
        assert os.path.exists(path)
        assert path.endswith('.docx')
        os.unlink(path)

    def test_header_table_date(self, generated_doc):
        t0 = generated_doc.tables[0]
        assert t0.rows[0].cells[0].text.strip() == 'Date:'

    def test_header_table_agreement_id(self, generated_doc, sample_data):
        t0 = generated_doc.tables[0]
        assert sample_data['AGREEMENT_ID'] in t0.rows[1].cells[1].text

    def test_header_table_description(self, generated_doc, sample_data):
        t0 = generated_doc.tables[0]
        assert sample_data['AGREEMENT_NAME_DESCRIPTION'] in t0.rows[2].cells[1].text

    def test_header_table_term(self, generated_doc, sample_data):
        t0 = generated_doc.tables[0]
        term_text = t0.rows[4].cells[1].text
        assert sample_data['START_DATE_YYYY_MM_DD'] in term_text
        assert sample_data['END_DATE_YYYY_MM_DD'] in term_text

    def test_financial_table_service(self, generated_doc, sample_data):
        t1 = generated_doc.tables[1]
        assert t1.rows[1].cells[0].text.strip() == sample_data['ITS_SERVICE']
        assert t1.rows[1].cells[1].text.strip() == sample_data['ITS_SERVICE_TYPE']
        assert t1.rows[1].cells[2].text.strip() == sample_data['AGREEMENT_TYPE']

    def test_financial_table_amount(self, generated_doc, sample_data):
        t1 = generated_doc.tables[1]
        assert t1.rows[1].cells[4].text.strip() == sample_data['MONTHLY_RECURRING']


class TestIFISCode:
    def test_ifis_fills_row0_not_row1(self, generated_doc):
        t2 = generated_doc.tables[2]
        # Row 0 should have data
        assert t2.rows[0].cells[0].text.strip() != ''
        # Row 1 should still have labels
        assert 'MINISTRY' in t2.rows[1].cells[0].text

    def test_ifis_separators_empty(self, generated_doc):
        t2 = generated_doc.tables[2]
        # Separator positions (3, 10, 15, 22) should be empty in data row
        for pos in [3, 10, 15, 22]:
            assert t2.rows[0].cells[pos].text.strip() == ''

    def test_ifis_correct_char_placement(self, generated_doc):
        t2 = generated_doc.tables[2]
        row = t2.rows[0]
        # IFIS_CODE: '100-123456-7890-ABCDEF-1234'
        # Stripped: '1001234567890ABCDEF1234' (23 chars)
        # Ministry (3): 1,0,0
        assert row.cells[0].text == '1'
        assert row.cells[1].text == '0'
        assert row.cells[2].text == '0'
        # Program Unit (6): 1,2,3,4,5,6
        assert row.cells[4].text == '1'
        assert row.cells[9].text == '6'
        # Business Unit (4): 7,8,9,0
        assert row.cells[11].text == '7'
        assert row.cells[14].text == '0'
        # Cost Centre (6): A,B,C,D,E,F
        assert row.cells[16].text == 'A'
        assert row.cells[21].text == 'F'
        # Initiative Code (4): 1,2,3,4
        assert row.cells[23].text == '1'
        assert row.cells[26].text == '4'

    def test_ifis_empty_code(self, sample_data):
        sample_data['IFIS_CODE'] = ''
        path = generate_docx.generate(sample_data)
        doc = Document(path)
        t2 = doc.tables[2]
        # All data cells should be empty
        for i, cell in enumerate(t2.rows[0].cells):
            assert cell.text.strip() == ''
        os.unlink(path)


class TestFooter:
    def test_prepared_by_contains_author(self, generated_doc, sample_data):
        found = False
        for section in generated_doc.sections:
            for para in section.footer.paragraphs:
                if 'Prepared By' in para.text:
                    assert sample_data['AGREEMENT_AUTHOR'] in para.text
                    found = True
        assert found, "No 'Prepared By' paragraph found in footer"

    def test_footer_rn_id(self, generated_doc, sample_data):
        found = False
        for section in generated_doc.sections:
            for para in section.footer.paragraphs:
                if 'RN ID' in para.text:
                    assert sample_data['AGREEMENT_ID'] in para.text
                    found = True
        assert found, "No 'RN ID' paragraph found in footer"
