"""Generate a filled Recovery Note Word document from submission data."""

from copy import deepcopy
import os
import re
import tempfile
from datetime import date

from docx import Document

# Path to the Word template — relative to this file's directory (tools/)
_TEMPLATE = os.path.normpath(
    os.path.join(
        os.path.dirname(os.path.abspath(__file__)), "..", "recovery_note_template.docx"
    )
)


def _para_full_text(para):
    """Return the concatenated text of all runs in a paragraph."""
    return "".join(r.text for r in para.runs)


def _set_para_text(para, text):
    """Replace all runs in a paragraph with a single run holding *text*."""
    text = text or ""
    if not para.runs:
        para.add_run(text)
        return
    para.runs[0].text = text
    for run in para.runs[1:]:
        run._r.getparent().remove(run._r)


def _fill_cell(cell, text):
    """Set the first paragraph of a table cell to *text*."""
    if cell.paragraphs:
        _set_para_text(cell.paragraphs[0], text or "")


def _copy_cell_style(src_cell, dst_cell):
    """Copy cell and paragraph/run formatting from *src_cell* to *dst_cell*."""
    src_tcpr = src_cell._tc.tcPr
    if src_tcpr is not None:
        dst_tcpr = dst_cell._tc.tcPr
        if dst_tcpr is not None:
            dst_cell._tc.remove(dst_tcpr)
        dst_cell._tc.append(deepcopy(src_tcpr))

    if not src_cell.paragraphs or not dst_cell.paragraphs:
        return

    src_para = src_cell.paragraphs[0]
    dst_para = dst_cell.paragraphs[0]
    dst_para.style = src_para.style
    dst_para.paragraph_format.alignment = src_para.paragraph_format.alignment
    dst_para.paragraph_format.left_indent = src_para.paragraph_format.left_indent
    dst_para.paragraph_format.right_indent = src_para.paragraph_format.right_indent
    dst_para.paragraph_format.first_line_indent = (
        src_para.paragraph_format.first_line_indent
    )
    dst_para.paragraph_format.space_before = src_para.paragraph_format.space_before
    dst_para.paragraph_format.space_after = src_para.paragraph_format.space_after
    dst_para.paragraph_format.line_spacing = src_para.paragraph_format.line_spacing

    if src_para.runs and dst_para.runs:
        src_font = src_para.runs[0].font
        dst_font = dst_para.runs[0].font
        dst_font.name = src_font.name
        dst_font.size = src_font.size
        dst_font.bold = src_font.bold
        dst_font.italic = src_font.italic
        dst_font.underline = src_font.underline


def _fill_ifis_row(row, code, label_row=None):
    """
    Fill the IFIS table row one character per data cell.
    Separator positions are determined from the label row (cells with '-').
    If no label_row is given, separators are detected from the data row itself.
    """
    # Determine separator positions from label row
    separator_positions = set()
    ref = label_row if label_row is not None else row
    for idx, cell in enumerate(ref.cells):
        if cell.text.strip() == "-":
            separator_positions.add(idx)

    chars = re.sub(r"[^0-9A-Za-z]", "", code) if code else ""
    char_idx = 0
    for idx, cell in enumerate(row.cells):
        if idx in separator_positions:
            continue
        _fill_cell(cell, chars[char_idx] if char_idx < len(chars) else "")
        char_idx += 1


def _format_currency(value):
    """Format numeric-like values as currency text for the document."""
    if value is None:
        return ""
    raw = str(value).strip()
    if not raw:
        return ""
    raw = raw.replace("$", "").replace(",", "").strip()
    try:
        num = float(raw)
    except ValueError:
        return f"${str(value).strip()}"
    return f"${num:.2f}"


def generate(data):
    """
    Fill the Recovery Note template with submission data.

    Parameters
    ----------
    data : dict
        Submission data keyed by sanitised field names (as stored in the DB).

    Returns
    -------
    str
        Absolute path to a temporary .docx file.  The caller is responsible
        for deleting it after the response has been sent.
    """
    d = data or {}
    doc = Document(_TEMPLATE)

    # ── Table 0: key header fields ────────────────────────────────────────────
    t0 = doc.tables[0]

    # Row 0: Date
    _fill_cell(t0.rows[0].cells[1], d.get("_created_at", date.today().isoformat()))

    # Row 1: Agreement ID
    _fill_cell(t0.rows[1].cells[1], d.get("AGREEMENT_ID", ""))

    # Row 2: Agreement Name / Description
    _fill_cell(t0.rows[2].cells[1], d.get("AGREEMENT_NAME_DESCRIPTION", ""))

    # Row 3: Previous Recovery Agreement
    _fill_cell(t0.rows[3].cells[1], d.get("PREVIOUS_AGREEMENT", ""))

    # Row 4: Term — start / end dates
    start = d.get("START_DATE_YYYY_MM_DD", "")
    end = d.get("END_DATE_YYYY_MM_DD", "")
    _fill_cell(t0.rows[4].cells[1], f"Start Date: {start}   End Date: {end}")

    # Row 5: Solution CI (always present)
    solution_ci = d.get("SOLUTION_CI", "").replace("\n", " ").strip()
    row5 = t0.add_row()

    # Match the new row to the existing header table formatting.
    _copy_cell_style(t0.rows[4].cells[0], row5.cells[0])
    _copy_cell_style(t0.rows[4].cells[1], row5.cells[1])

    _fill_cell(row5.cells[0], "Solution CI:")
    _fill_cell(row5.cells[1], solution_ci)

    # ── Paragraph "Column R" → Comments ──────────────────────────────────────
    comments = d.get("COMMENTS", "")
    for para in doc.paragraphs:
        if "Column R" in _para_full_text(para):
            _set_para_text(para, comments)
            break

    # ── Table 1: financial recovery details ───────────────────────────────────
    t1 = doc.tables[1]
    row1 = t1.rows[1]
    _fill_cell(row1.cells[0], d.get("ITS_SERVICE", "").replace("\n", " ").strip())
    _fill_cell(row1.cells[1], d.get("ITS_SERVICE_TYPE", "").replace("\n", " ").strip())
    _fill_cell(row1.cells[2], d.get("AGREEMENT_TYPE", ""))
    _fill_cell(row1.cells[3], d.get("MONTH_BILLED", ""))
    amount = d.get("MONTHLY_RECURRING") or d.get("ANNUAL") or d.get("ONE_TIME") or ""
    _fill_cell(row1.cells[4], _format_currency(amount))

    # ── Table 2: IFIS breakdown ───────────────────────────────────────────────
    t2 = doc.tables[2]
    _fill_ifis_row(t2.rows[0], d.get("IFIS_CODE", ""), label_row=t2.rows[1])

    # ── Paragraph "Column X" → full IFIS code ─────────────────────────────────
    for para in doc.paragraphs:
        if "Column X" in _para_full_text(para):
            _set_para_text(para, d.get("IFIS_CODE", ""))
            break

    # ── Cluster / approver section ─────────────────────────────────────────────
    cluster = d.get("SERVICE_OWNER", "")
    for para in doc.paragraphs:
        if _para_full_text(para).strip() == "Name of Cluster":
            _set_para_text(para, cluster)
            break

    # ── Footer: RN ID and Prepared By ─────────────────────────────────────────
    agreement_id = d.get("AGREEMENT_ID", "")
    author = d.get("AGREEMENT_AUTHOR", "")
    for section in doc.sections:
        for para in section.footer.paragraphs:
            full = _para_full_text(para)
            if "Prepared By" in full:
                _set_para_text(para, f"Prepared By: {author}")
            elif "RN ID" in full:
                _set_para_text(para, f"RN ID:  {agreement_id}")

    # ── Write to a temporary file ──────────────────────────────────────────────
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()
    doc.save(tmp.name)
    return tmp.name
